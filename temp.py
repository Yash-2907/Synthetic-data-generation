import os
import random
import requests
import wikipedia
from faker import Faker
from docx import Document
from PyPDF2 import PdfReader, PdfWriter
from io import BytesIO
import csv
import time
import logging

logger = logging.getLogger(__name__)
logger.setLevel(logging.INFO)

c_handler = logging.StreamHandler()
c_handler.setLevel(logging.INFO)

log_file_path = "script_activity.log"
f_handler = logging.FileHandler(log_file_path, mode='a')
f_handler.setLevel(logging.DEBUG)

c_format = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
f_format = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(funcName)s - %(message)s')
c_handler.setFormatter(c_format)
f_handler.setFormatter(f_format)

logger.addHandler(c_handler)
logger.addHandler(f_handler)

fake = Faker(['en_US', 'th_TH', 'id_ID'])

NUM_FILES_PER_TYPE = 100

base_dir = "sample_files"
types = ["txt", "docx", "csv", "pdf", "modified"]
for t in types:
    os.makedirs(f"{base_dir}/{t}", exist_ok=True)
    if t != "modified":
        os.makedirs(f"{base_dir}/modified/{t}", exist_ok=True)
logger.info(f"Ensured base directories '{base_dir}' and subdirectories exist.")

email_domains = [
    "gmail.com", "yahoo.com", "hotmail.com", "outlook.com", "live.com",
    "aol.com", "protonmail.com", "icloud.com", "mail.com", "zoho.com",
    "company.com", "org.net" # Added a few more generic/business domains
]

base_fields = [
    "Customer_ID", "Full_Name", "Gender", "Date_of_Birth", "Nationality",
    "Account_Number", "Bank_Name", "SWIFT_Code", "IBAN", "Account_Type",
    "Opening_Date", "Current_Balance", "Currency",
    "Loan_ID", "Loan_Amount", "Loan_Purpose", "Loan_Status", "Collateral_Type",
    "Employment_Status", "Monthly_Income", "Monthly_Expenses", "Employer_Name", "Job_Title",
    "MAC_Address", "IP_Address", "Device_Type", "Login_Timestamp", "Logout_Timestamp",
    "Account_Locked", "Two_Factor_Enabled", "Biometric_Auth_Enabled",
    "Crypto_Wallet_Address", "Crypto_Balance", "Wallet_Type",
    "Reward_Points", "Reward_Redemption_Rate", "Loyalty_Level",
    "Online_Shopping_Spend", "Utility_Bill_Payments", "Subscription_Services",
    "Support_Tickets", "Customer_Feedback", "Satisfaction_Score", "Complaint_Status",
    "Transaction_ID", "Transaction_Type", "Transaction_Amount", "Transaction_Date",
    "Transaction_Channel", "Merchant_Category", "Frequent_Merchant", "Top_Spending_Category"
]

sensitive_fields = [
    "Email", "Phone_Number", "TIN", "Thai_ID", "Passport_Number",
    "Social_Security_Number", "Credit_Card_Number", "Debit_Card_Number",
    "Bank_Account_Number", "Routing_Number", "Medical_Record_Number",
    "Driver_License_Number", "Biometric_Data_Fingerprint", "Biometric_Data_FaceScan",
    "Mother_Maiden_Name", "Security_Youtube", "PIN_Code", "Password_Hash",
    "Digital_Signature", "Private_Key", "Cryptographic_Seed", "API_Key",
    "Physical_Address", "GPS_Coordinates", "IP_Address_Exact", "Device_IMEI",
    "Vehicle_Identification_Number", "Patient_ID", "Health_Insurance_ID",
    "Employee_ID", "Payroll_ID", "Tax_ID", "Voter_ID",
    "Student_ID", "Library_Card_Number", "Membership_ID",
    "Genetic_Sequence_Data", "Blood_Type", "Allergies", "Diagnosis_Code",
    "Prescription_Details", "Insurance_Policy_Number", "Financial_Statement_Details",
    "Investment_Portfolio_Details", "Loan_Application_Details", "Credit_Score",
    "Criminal_Record_Details", "Litigation_Status", "Trade_Secret_Info",
    "Confidential_Project_Name", "Internal_Audit_Report", "Proprietary_Algorithm_Details",
    "Network_Credentials", "System_Access_Key", "Vendor_Proprietary_Info",
    "Customer_Purchase_History", "Browser_History", "Search_Query_Data",
    "Online_Activity_Log", "Communication_Content_Snippet"
]

def generate_sensitive_field():
    field = random.choice(sensitive_fields)
    if field == "Email":
        return f"Email: {fake.user_name()}@{random.choice(email_domains)}"
    elif field == "Phone_Number":
        return f"Phone_Number: {fake.phone_number()}"
    elif field == "TIN": # Assuming TIN is a US SSN for faker
        return f"TIN: {fake.ssn()}"
    elif field == "Thai_ID":
        # Faker might not have a direct Thai ID, using bothify as placeholder
        return f"Thai_ID: {fake.bothify(text='#############'[:13])}" 
    elif field == "Passport_Number":
        return f"Passport_Number: {fake.bothify(text='??########')}"
    elif field == "Social_Security_Number":
        return f"SSN: {fake.ssn()}"
    elif field == "Credit_Card_Number":
        return f"Credit_Card: {fake.credit_card_number()}"
    elif field == "Debit_Card_Number":
        return f"Debit_Card: {fake.credit_card_number(card_type='mastercard')}" # Faker doesn't distinguish debit, so use common card type
    elif field == "Bank_Account_Number":
        return f"Bank_Account: {fake.bban()}" # Basic Bank Account Number
    elif field == "Routing_Number":
        return f"Routing_Number: {fake.aba()}" # ABA Routing Number
    elif field == "Medical_Record_Number":
        return f"MRN: {fake.bothify(text='???######')}"
    elif field == "Driver_License_Number":
        return f"DLN: {fake.license_plate()}" # Using license plate as a generic ID faker
    elif field == "Biometric_Data_Fingerprint":
        return f"Biometric: {fake.sha256()[:16]}" # Placeholder for hash of biometric data
    elif field == "Biometric_Data_FaceScan":
        return f"Biometric: {fake.md5()[:16]}" # Placeholder for hash of biometric data
    elif field == "Mother_Maiden_Name":
        return f"Mother's Maiden Name: {fake.last_name()}"
    elif field == "Security_Youtube":
        return f"Security Answer: {fake.word()}"
    elif field == "PIN_Code":
        return f"PIN: {fake.bothify(text='####')}"
    elif field == "Password_Hash":
        return f"Password Hash: {fake.sha256()}"
    elif field == "Digital_Signature":
        return f"Digital Signature: {fake.sha512()}"
    elif field == "Private_Key":
        return f"Private Key: {fake.sha256()}{fake.sha256()}" # Longer random string
    elif field == "Cryptographic_Seed":
        return f"Crypto Seed: {fake.uuid4()}"
    elif field == "API_Key":
        return f"API Key: {fake.uuid4()}"
    elif field == "Physical_Address":
        return f"Address: {fake.address()}"
    elif field == "GPS_Coordinates":
        return f"GPS: Lat {fake.latitude()}, Long {fake.longitude()}"
    elif field == "IP_Address_Exact":
        return f"IP: {fake.ipv4()}"
    elif field == "Device_IMEI":
        return f"IMEI: {fake.msisdn()[:15]}" # MSISDN is phone number, truncating for IMEI length
    elif field == "Vehicle_Identification_Number":
        return f"VIN: {fake.vin()}"
    elif field == "Patient_ID":
        return f"Patient ID: {fake.bothify(text='P######')}"
    elif field == "Health_Insurance_ID":
        return f"Insurance ID: {fake.bothify(text='HIC##########')}"
    elif field == "Employee_ID":
        return f"Employee ID: {fake.bothify(text='EMP###')}"
    elif field == "Payroll_ID":
        return f"Payroll ID: {fake.bothify(text='PAY####')}"
    elif field == "Tax_ID": # Generic tax ID
        return f"Tax ID: {fake.bothify(text='########')}"
    elif field == "Voter_ID":
        return f"Voter ID: {fake.bothify(text='VOT###')}"
    elif field == "Student_ID":
        return f"Student ID: {fake.bothify(text='STD#####')}"
    elif field == "Library_Card_Number":
        return f"Library Card: {fake.bothify(text='LC######')}"
    elif field == "Membership_ID":
        return f"Membership ID: {fake.bothify(text='MEM######')}"
    elif field == "Genetic_Sequence_Data":
        return f"Genetic Data: {fake.sha256()[:20]}..." # Placeholder for complex data
    elif field == "Blood_Type":
        return f"Blood Type: {random.choice(['A+', 'A-', 'B+', 'B-', 'AB+', 'AB-', 'O+', 'O-'])}"
    elif field == "Allergies":
        return f"Allergies: {fake.word()} pollen, {fake.word()} nuts"
    elif field == "Diagnosis_Code":
        return f"Diagnosis: ICD-{fake.bothify(text='???####')}"
    elif field == "Prescription_Details":
        return f"Prescription: {fake.word()} {fake.bothify(text='##')}"
    elif field == "Insurance_Policy_Number":
        return f"Policy No: {fake.bothify(text='POL#####')}"
    elif field == "Financial_Statement_Details":
        return f"Financial Detail: {fake.sentence()}"
    elif field == "Investment_Portfolio_Details":
        return f"Investment Detail: {fake.sentence()}"
    elif field == "Loan_Application_Details":
        return f"Loan App: {fake.sentence()}"
    elif field == "Credit_Score":
        return f"Credit Score: {random.randint(300, 850)}"
    elif field == "Criminal_Record_Details":
        return f"Criminal Record: {fake.sentence()}"
    elif field == "Litigation_Status":
        return f"Litigation: {random.choice(['Pending', 'Closed', 'Settled'])}"
    elif field == "Trade_Secret_Info":
        return f"Trade Secret: {fake.sentence()}"
    elif field == "Confidential_Project_Name":
        return f"Confidential Project: {fake.catch_phrase()}"
    elif field == "Internal_Audit_Report":
        return f"Audit Report: {fake.sentence()}"
    elif field == "Proprietary_Algorithm_Details":
        return f"Algorithm: {fake.sentence()}"
    elif field == "Network_Credentials":
        return f"Net Creds: User:{fake.user_name()} Pass:{fake.password()}"
    elif field == "System_Access_Key":
        return f"System Key: {fake.uuid4()}"
    elif field == "Vendor_Proprietary_Info":
        return f"Vendor Info: {fake.sentence()}"
    elif field == "Customer_Purchase_History":
        return f"Purchase History: {fake.sentence()}"
    elif field == "Browser_History":
        return f"Browser History: {fake.url()}"
    elif field == "Search_Query_Data":
        return f"Search Query: {fake.sentence()}"
    elif field == "Online_Activity_Log":
        return f"Activity Log: {fake.sentence()}"
    elif field == "Communication_Content_Snippet":
        return f"Communication: {fake.sentence()}"
    
    return ""

def generate_base_field():
    return f"{random.choice(base_fields)}: {fake.word()}"

def download_texts():
    logger.info("--- Downloading TXT files from Project Gutenberg ---")
    successful_downloads = 0
    tested_ids = set()
    
    potential_ids = list(range(1, 70000))
    random.shuffle(potential_ids)

    for book_id in potential_ids:
        if successful_downloads >= NUM_FILES_PER_TYPE:
            break
        if book_id in tested_ids:
            continue
        tested_ids.add(book_id)

        url = f"https://www.gutenberg.org/files/{book_id}/{book_id}-0.txt"
        
        try:
            r = requests.get(url, timeout=10)
            if r.ok:
                if "text/plain" in r.headers.get("Content-Type", ""):
                    path = f"{base_dir}/txt/book_{book_id}.txt"
                    with open(path, "w", encoding="utf-8") as f:
                        f.write(r.text)
                    logger.info(f"[TXT] Downloaded book_{book_id}.txt ({successful_downloads + 1}/{NUM_FILES_PER_TYPE})")
                    successful_downloads += 1
                else:
                    logger.debug(f"[TXT] Skipping {url}: Not a plain text file (Content-Type: {r.headers.get('Content-Type', 'N/A')}).")
            else:
                logger.warning(f"[TXT] Failed to download {url} (Status: {r.status_code}). Moving to next.")
        except requests.exceptions.RequestException as e:
            logger.error(f"[TXT] Error downloading {url}: {e}. Moving to next.")
    logger.info(f"Finished downloading TXT files. Total: {successful_downloads}")

def create_docx():
    logger.info("--- Creating DOCX files from Wikipedia ---")
    num_created = 0
    while num_created < NUM_FILES_PER_TYPE:
        try:
            page_title = wikipedia.random(1)
            content = wikipedia.page(page_title, auto_suggest=False, redirect=True).content
            
            doc = Document()
            doc.add_heading(page_title, 0)
            doc.add_paragraph(content[:8000])
            
            filename = "".join([c if c.isalnum() or c in [' ', '_', '-'] else '_' for c in page_title]).replace(' ', '_')
            if not filename:
                filename = f"random_doc_{num_created}"
            
            path = f"{base_dir}/docx/{filename}.docx"
            
            counter = 0
            original_path = path
            while os.path.exists(path):
                counter += 1
                path = f"{original_path.rsplit('.', 1)[0]}_{counter}.docx"

            doc.save(path)
            logger.info(f"[DOCX] Created {os.path.basename(path)} ({num_created + 1}/{NUM_FILES_PER_TYPE})")
            num_created += 1
        except wikipedia.exceptions.DisambiguationError as e:
            logger.warning(f"[DOCX] Skipping disambiguation page '{e.title}': Options: {', '.join(e.options[:3])}...")
        except wikipedia.exceptions.PageError:
            logger.warning(f"[DOCX] Skipping page not found for '{page_title}'.")
        except requests.exceptions.RequestException as e:
            logger.error(f"[DOCX] Network error fetching Wikipedia page '{page_title}': {e}.")
        except Exception as e:
            logger.error(f"[DOCX] Generic error creating DOCX for '{page_title}': {e}", exc_info=True)
        time.sleep(0.05)
    logger.info(f"Finished creating DOCX files. Total: {num_created}")


def download_csv():
    logger.info("--- Downloading/Creating CSV files ---")
    # Expanded list of public CSV URLs for download
    urls = []
    
    num_downloaded = 0
    for i, url in enumerate(urls):
        if num_downloaded >= NUM_FILES_PER_TYPE:
            break
        try:
            r = requests.get(url, timeout=10)
            if r.ok:
                path = f"{base_dir}/csv/downloaded_sample_{i+1}.csv"
                with open(path, "wb") as f:
                    f.write(r.content)
                logger.info(f"[CSV] Downloaded {os.path.basename(path)}")
                num_downloaded += 1
            else:
                logger.warning(f"[CSV] Failed to download {url} (Status: {r.status_code}). Moving to next.")
        except requests.exceptions.RequestException as e:
            logger.error(f"[CSV] Error downloading {url}: {e}. Moving to next.")

    while num_downloaded < NUM_FILES_PER_TYPE:
        filename = f"{base_dir}/csv/synthetic_data_{num_downloaded + 1}.csv"
        try:
            with open(filename, 'w', newline='', encoding='utf-8') as csvfile:
                writer = csv.writer(csvfile)
                writer.writerow(['ID', 'Name', 'Age', 'City', 'Product', 'Quantity', 'Price'])
                for _ in range(random.randint(50, 100)):
                    writer.writerow([
                        fake.uuid4(),
                        fake.name(),
                        fake.random_int(min=18, max=80),
                        fake.city(),
                        fake.word().capitalize(),
                        fake.random_int(min=1, max=10),
                        round(fake.random_int(min=10, max=1000) / 100.0, 2)
                    ])
            logger.info(f"[CSV] Created {os.path.basename(filename)} ({num_downloaded + 1}/{NUM_FILES_PER_TYPE})")
            num_downloaded += 1
        except Exception as e:
            logger.error(f"[CSV] Error creating synthetic CSV {filename}: {e}", exc_info=True)
    logger.info(f"Finished downloading/creating CSV files. Total: {num_downloaded}")

def download_pdfs():
    logger.info("--- Downloading PDF files ---")
    # List of real PDF URLs for testing
    urls = []
    
    num_downloaded = 0
    for i, url in enumerate(urls):
        if num_downloaded >= NUM_FILES_PER_TYPE:
            break
        try:
            r = requests.get(url, stream=True, timeout=15)
            if r.ok:
                path = f"{base_dir}/pdf/downloaded_paper_{i+1}.pdf"
                with open(path, "wb") as f:
                    for chunk in r.iter_content(chunk_size=8192):
                        f.write(chunk)
                logger.info(f"[PDF] Downloaded {os.path.basename(path)}")
                num_downloaded += 1
            else:
                logger.warning(f"[PDF] Failed to download {url} (Status: {r.status_code}). Moving to next.")
        except requests.exceptions.RequestException as e:
            logger.error(f"[PDF] Error downloading {url}: {e}. Moving to next.")

    while num_downloaded < NUM_FILES_PER_TYPE:
        filename = f"{base_dir}/pdf/empty_pdf_{num_downloaded + 1}.pdf"
        try:
            writer = PdfWriter()
            writer.add_blank_page(width=72 * 8.5, height=72 * 11)
            with open(filename, "wb") as f_out:
                writer.write(f_out)
            logger.info(f"[PDF] Created {os.path.basename(filename)} (empty) ({num_downloaded + 1}/{NUM_FILES_PER_TYPE})")
            num_downloaded += 1
        except Exception as e:
            logger.error(f"[PDF] Error creating empty PDF {filename}: {e}", exc_info=True)
    logger.info(f"Finished downloading/creating PDF files. Total: {num_downloaded}")


def inject_txt():
    logger.info("\n--- Injecting into TXT files ---")
    files_to_process = os.listdir(f"{base_dir}/txt")
    if not files_to_process:
        logger.warning("No TXT files found to inject into.")
        return
    for f in files_to_process:
        src = os.path.join(base_dir, "txt", f)
        dst = os.path.join(base_dir, "modified/txt", f)
        try:
            lines = open(src, "r", encoding="utf-8", errors='ignore').readlines()
            insert = generate_sensitive_field() if random.random() < 0.5 else generate_base_field()
            idx = random.randint(min(1, len(lines)), max(1, len(lines)-1)) if len(lines) > 0 else 0
            lines.insert(idx, insert + "\n")
            with open(dst, "w", encoding="utf-8") as f_out:
                f_out.writelines(lines)
            logger.info(f"[TXT] Injected into {f} at line {idx}")
        except Exception as e:
            logger.error(f"[TXT] Error injecting into {f}: {e}", exc_info=True)
    logger.info("Finished injecting into TXT files.")

def inject_docx():
    logger.info("\n--- Injecting into DOCX files ---")
    files_to_process = os.listdir(f"{base_dir}/docx")
    if not files_to_process:
        logger.warning("No DOCX files found to inject into.")
        return

    for f in files_to_process:
        src = os.path.join(base_dir, "docx", f)
        dst = os.path.join(base_dir, "modified/docx", f)
        try:
            doc = Document(src)
            num_paragraphs = len(doc.paragraphs)

            num_sensitive_fields_to_inject = random.randint(0, 4)

            for _ in range(num_sensitive_fields_to_inject):
                injected_data = generate_sensitive_field()

                if num_paragraphs == 0:
                    doc.add_paragraph(injected_data)
                    logger.info(f"[DOCX] Injected into empty {f}.")
                else:
                    injection_type = random.choices(
                        ["new_paragraph", "modify_existing_paragraph"],
                        weights=[0.6, 0.4],
                        k=1
                    )[0]

                    if injection_type == "new_paragraph":
                        if num_paragraphs == 1:
                            insert_idx = random.choice([0, 1])
                        elif num_paragraphs > 1:
                            possible_indices = list(range(0, num_paragraphs + 1))
                            weights = [1] * len(possible_indices)
                            weights[0] = 0.5
                            weights[num_paragraphs] = 0.5
                            insert_idx = random.choices(possible_indices, weights=weights, k=1)[0]
                        
                        if insert_idx == num_paragraphs:
                            doc.add_paragraph(injected_data)
                            logger.info(f"[DOCX] Injected into {f} at the very end (new paragraph).")
                        else:
                            target_paragraph = doc.paragraphs[insert_idx]
                            target_paragraph.insert_paragraph_before(injected_data)
                            logger.info(f"[DOCX] Injected into {f} before paragraph {insert_idx} (new paragraph).")

                    else:
                        target_paragraph_idx = random.randint(0, num_paragraphs - 1)
                        target_paragraph = doc.paragraphs[target_paragraph_idx]
                        
                        target_paragraph.add_run(f" {injected_data}")
                        logger.info(f"[DOCX] Appended new run to paragraph {target_paragraph_idx} in {f}.")
                
                num_paragraphs = len(doc.paragraphs)

            doc.save(dst)
        except Exception as e:
            logger.error(f"[DOCX] Error injecting into {f}: {e}", exc_info=True)
    logger.info("Finished injecting into DOCX files.")

def inject_csv():
    logger.info("\n--- Injecting into CSV files ---")
    files_to_process = os.listdir(f"{base_dir}/csv")
    if not files_to_process:
        logger.warning("No CSV files found to inject into.")
        return
    for f in files_to_process:
        src = os.path.join(base_dir, "csv", f)
        dst = os.path.join(base_dir, "modified/csv", f)
        try:
            lines = open(src, "r", encoding="utf-8", errors='ignore').readlines()
            insert_data = generate_sensitive_field() if random.random() < 0.5 else generate_base_field()
            idx = random.randint(1, len(lines)-1) if len(lines) > 1 else 1
            lines.insert(idx, f"{insert_data}\n")
            with open(dst, "w", encoding="utf-8") as f_out:
                f_out.writelines(lines)
            logger.info(f"[CSV] Injected into {f} at line {idx}")
        except Exception as e:
            logger.error(f"[CSV] Error injecting into {f}: {e}", exc_info=True)
    logger.info("Finished injecting into CSV files.")

def inject_pdf():
    logger.info("\n--- Injecting into PDF files ---")
    files_to_process = os.listdir(f"{base_dir}/pdf")
    if not files_to_process:
        logger.warning("No PDF files found to inject into.")
        return
    for f in files_to_process:
        src = os.path.join(base_dir, "pdf", f)
        dst = os.path.join(base_dir, "modified/pdf", f)
        try:
            reader = PdfReader(src)
            writer = PdfWriter()
            for page in reader.pages:
                writer.add_page(page)
            
            meta_key = random.choice(["Author", "Subject", "Title", "Keywords", "Creator"])
            meta_value = generate_sensitive_field() if random.random() < 0.5 else generate_base_field()
            
            existing_meta = reader.metadata or {}
            new_meta = {f"/{k}": v for k, v in existing_meta.items()}
            new_meta[f"/{meta_key}"] = meta_value

            writer.add_metadata(new_meta)
            
            with open(dst, "wb") as f_out:
                writer.write(f_out)
            logger.info(f"[PDF] Injected '{meta_key}' with data into {f}")
        except Exception as e:
            logger.error(f"[PDF] Skipped {f} due to error during injection: {e}", exc_info=True)
    logger.info("Finished injecting into PDF files.")

def main():
    logger.info("--- Starting File Generation and Injection Process ---")
    logger.info(f"Target: {NUM_FILES_PER_TYPE} files per type.")

    logger.info("\n--- Beginning File Downloads/Creations ---")
    #download_texts()
    #create_docx()
    #download_csv()
    #download_pdfs()

    logger.info("\n--- Initiating Data Injection ---")
    #inject_txt()
    inject_docx()
    #inject_csv()
    #inject_pdf()
    
    logger.info("\n--- ✅ All Done. Check 'sample_files' directory and 'script_activity.log' for details. ---")

if __name__ == "__main__":
    main()