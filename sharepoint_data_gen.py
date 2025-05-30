import os
import requests
from docx import Document
import wikipedia

# Create folders
os.makedirs("sample_files/txt", exist_ok=True)
os.makedirs("sample_files/docx", exist_ok=True)
os.makedirs("sample_files/csv", exist_ok=True)
os.makedirs("sample_files/xlsx", exist_ok=True)
os.makedirs("sample_files/pdf", exist_ok=True)

# ---- TEXT FILES from Gutenberg ----
def download_gutenberg_texts():
    ids = [1342, 1661]  # Pride and Prejudice, Sherlock Holmes
    for book_id in ids:
        url = f"https://www.gutenberg.org/files/{book_id}/{book_id}-0.txt"
        response = requests.get(url)
        if response.status_code == 200:
            with open(f"sample_files/txt/book_{book_id}.txt", "w", encoding="utf-8") as f:
                f.write(response.text)

# ---- DOCX FILES from Wikipedia ----
def create_docx_from_wikipedia():
    topics = ["Artificial intelligence", "Cloud computing"]
    for topic in topics:
        try:
            content = wikipedia.page(topic).content
            doc = Document()
            doc.add_heading(topic, 0)
            doc.add_paragraph(content[:5000])
            safe_name = topic.replace(" ", "_").replace("/", "_")
            doc.save(f"sample_files/docx/{safe_name}.docx")
        except:
            continue

# ---- CSV FILES from DataHub.io ----
def download_csv_files():
    urls = [
        "https://datahub.io/core/global-temp/r/annual.csv",
        "https://datahub.io/core/co2-fossil-global/r/global.csv"
    ]
    for i, url in enumerate(urls):
        response = requests.get(url)
        with open(f"sample_files/csv/sample_{i+1}.csv", "wb") as f:
            f.write(response.content)

# ---- XLSX FILES from World Bank ----
def download_xlsx_files():
    urls = [
        "https://databankfiles.worldbank.org/public/ddpext_download/DoingBusiness/DB19/XLS/Dataset_CHN.xlsx",
        "https://databankfiles.worldbank.org/public/ddpext_download/DoingBusiness/DB19/XLS/Dataset_USA.xlsx"
    ]
    for i, url in enumerate(urls):
        response = requests.get(url)
        with open(f"sample_files/xlsx/sample_{i+1}.xlsx", "wb") as f:
            f.write(response.content)

# ---- PDF FILES from arXiv ----
def download_pdf_files():
    urls = [
        "https://arxiv.org/pdf/2106.14806.pdf",
        "https://arxiv.org/pdf/2301.10226.pdf"
    ]
    for i, url in enumerate(urls):
        response = requests.get(url)
        with open(f"sample_files/pdf/paper_{i+1}.pdf", "wb") as f:
            f.write(response.content)

# Run all
download_gutenberg_texts()
create_docx_from_wikipedia()
download_csv_files()
download_xlsx_files()
download_pdf_files()

print("✅ Files downloaded into 'sample_files' folder.")
